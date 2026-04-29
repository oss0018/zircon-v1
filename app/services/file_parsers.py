"""
Text extractors for supported file types.
"""
import json
import csv
import io
import re
import logging
from pathlib import Path
from typing import Optional

MAX_INDEX_BYTES = 50 * 1024 * 1024  # 50 MB max content to index

logger = logging.getLogger(__name__)

# Encoding fallback order for plain-text files
_ENCODINGS = ("utf-8", "cp1251", "latin-1")


def _read_text_with_encoding(file_path: str, max_bytes: int = MAX_INDEX_BYTES) -> str:
    """Read a text file trying UTF-8, cp1251, then latin-1.

    Optionally auto-detect encoding via chardet when available.
    """
    raw: Optional[bytes] = None
    try:
        with open(file_path, "rb") as f:
            raw = f.read(max_bytes)
    except Exception as e:
        logger.warning("[file_parsers] Cannot read %s: %s", file_path, e)
        return ""

    # Try chardet first if available
    try:
        import chardet
        detected = chardet.detect(raw)
        enc = detected.get("encoding") or "utf-8"
        return raw.decode(enc, errors="replace")
    except ImportError:
        pass

    for enc in _ENCODINGS:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _is_credential_log(p: Path) -> bool:
    # Check if file is inside a leaked_accounts directory
    parts = p.parts
    if "leaked_accounts" in parts:
        return True
    # Sample first few lines to detect the pattern
    if p.suffix.lower() in (".txt", ".log", ""):
        try:
            with open(p, "r", errors="replace") as f:
                for _ in range(10):
                    line = f.readline()
                    if not line:
                        break
                    line = line.strip()
                    if not line or line.startswith("#"):
                        continue
                    # Match url:login:pass or email:password patterns
                    if re.match(r'^https?://[^\s:]+[:].+[:].+$', line):
                        return True
                    if re.match(r'^[\w.+-]+@[\w-]+\.[\w.]+[:].+$', line):
                        return True
        except Exception:
            pass
    return False


def extract_credential_log(file_path: str, max_bytes: int = MAX_INDEX_BYTES) -> str:
    """Parse url:login:password logs, extract searchable tokens."""
    domains: set = set()
    emails: set = set()
    usernames: set = set()
    raw_lines = []

    bytes_read = 0
    try:
        with open(file_path, "r", errors="replace") as f:
            for line in f:
                bytes_read += len(line.encode("utf-8", errors="replace"))
                if bytes_read > max_bytes:
                    break
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                raw_lines.append(line)
                # Reconstruct URL:login:pass — handle http:// prefix
                raw_parts = line.split(":")
                if len(raw_parts) >= 3 and raw_parts[0].lower() in ("http", "https"):
                    # URL format: https://domain/path:login:pass
                    url_part = f"{raw_parts[0]}:{raw_parts[1]}"
                    rest = raw_parts[2:]
                else:
                    url_part = raw_parts[0]
                    rest = raw_parts[1:]

                # Extract domain from URL
                m = re.search(r"(?:https?://)?([^/:]+\.[^/:]+)", url_part)
                if m:
                    domains.add(m.group(1).lower())
                elif "." in url_part and "@" not in url_part:
                    domains.add(url_part.lower())

                # Extract emails and usernames from all parts
                for part in [url_part] + rest:
                    em = re.search(r"[\w.+-]+@[\w-]+\.[\w.]+", part)
                    if em:
                        emails.add(em.group(0).lower())
                        usernames.add(em.group(0).split("@")[0].lower())
    except Exception as e:
        return f"[credential log error: {e}]"

    parts_out = []
    parts_out.append("=== RAW LINES (sample) ===")
    parts_out.extend(raw_lines[:10000])
    parts_out.append("\n=== EXTRACTED DOMAINS ===")
    parts_out.extend(sorted(domains))
    parts_out.append("\n=== EXTRACTED EMAILS ===")
    parts_out.extend(sorted(emails))
    parts_out.append("\n=== EXTRACTED USERNAMES ===")
    parts_out.extend(sorted(usernames))

    return "\n".join(parts_out)


def extract_text_streaming(file_path: str, max_bytes: int = MAX_INDEX_BYTES) -> Optional[str]:
    """Read large files in chunks: first + middle + last chunk."""
    p = Path(file_path)
    if not p.exists():
        return None
    file_size = p.stat().st_size

    if file_size <= max_bytes:
        return extract_text(file_path)

    chunks = []
    chunk_size = max_bytes // 3

    try:
        with open(file_path, "rb") as f:
            chunks.append(f.read(chunk_size).decode("utf-8", errors="replace"))
            f.seek(file_size // 2)
            chunks.append(f.read(chunk_size).decode("utf-8", errors="replace"))
            f.seek(max(0, file_size - chunk_size))
            chunks.append(f.read(chunk_size).decode("utf-8", errors="replace"))
    except Exception as e:
        return f"[large file read error: {e}]"

    size_gb = file_size / (1024 ** 3)
    header = f"[LARGE FILE: {size_gb:.2f} GB, partial index]\n"
    return header + "\n...[TRUNCATED]...\n".join(chunks)


def extract_text(file_path: str) -> Optional[str]:
    p = Path(file_path)
    if not p.exists():
        return None

    # Check for large files first
    try:
        file_size = p.stat().st_size
    except Exception:
        file_size = 0

    # Detect credential log files and use specialized parser
    if _is_credential_log(p):
        return extract_credential_log(file_path)

    # For large plain-text files, use streaming extraction
    if file_size > MAX_INDEX_BYTES:
        return extract_text_streaming(file_path)

    suffix = p.suffix.lower()
    try:
        if suffix in (".txt", ".md", ".log", ".sql", ".html", ".htm", ".yaml", ".yml", ".ini", ".cfg", ".conf"):
            return _read_text_with_encoding(file_path)
        elif suffix == ".xml":
            return _extract_xml(p)
        elif suffix == ".json":
            return _extract_json(p)
        elif suffix == ".csv":
            return _extract_csv(p)
        elif suffix == ".xlsx":
            return _extract_excel(p)
        elif suffix == ".xls":
            return _extract_xls(p)
        elif suffix == ".pdf":
            return _extract_pdf(p)
        elif suffix == ".docx":
            return _extract_docx(p)
        elif suffix == ".doc":
            return _extract_doc(p)
        elif suffix == ".odt":
            return _extract_odt(p)
        elif suffix == ".rtf":
            return _extract_rtf(p)
        else:
            # Try reading as plain text with encoding detection
            try:
                return _read_text_with_encoding(file_path)
            except Exception:
                return None
    except Exception as e:
        return f"[extraction error: {e}]"


def _extract_json(p: Path) -> str:
    try:
        text = _read_text_with_encoding(str(p))
        data = json.loads(text)
        return json.dumps(data, ensure_ascii=False, indent=2)
    except Exception:
        return _read_text_with_encoding(str(p))


def _extract_csv(p: Path) -> str:
    rows = []
    # Try encoding detection
    text = _read_text_with_encoding(str(p))
    try:
        reader = csv.reader(io.StringIO(text))
        for row in reader:
            rows.append(" ".join(row))
    except Exception:
        pass
    return "\n".join(rows)


def _extract_excel(p: Path) -> str:
    try:
        import pandas as pd
        dfs = pd.read_excel(p, sheet_name=None, engine="openpyxl")
        parts = []
        for sheet_name, df in dfs.items():
            parts.append(f"[Sheet: {sheet_name}]")
            parts.append(df.to_string(index=False))
        return "\n".join(parts)
    except Exception as e:
        logger.warning("[file_parsers] Excel (xlsx) error for %s: %s", p, e)
        return f"[excel error: {e}]"


def _extract_xls(p: Path) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(str(p))
        parts = []
        for sheet in wb.sheets():
            parts.append(f"[Sheet: {sheet.name}]")
            for row_idx in range(sheet.nrows):
                row_cells = [str(sheet.cell_value(row_idx, c)) for c in range(sheet.ncols)]
                parts.append(" ".join(row_cells))
        return "\n".join(parts)
    except ImportError:
        logger.warning("[file_parsers] xlrd not installed; cannot parse .xls")
        return ""
    except Exception as e:
        logger.warning("[file_parsers] XLS error for %s: %s", p, e)
        return f"[xls error: {e}]"


def _extract_xml(p: Path) -> str:
    try:
        import xml.etree.ElementTree as ET
        text = _read_text_with_encoding(str(p))
        root = ET.fromstring(text)
        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            if elem.tail and elem.tail.strip():
                texts.append(elem.tail.strip())
        return "\n".join(texts)
    except Exception:
        # Fall back to plain text read (strips tags roughly)
        return _read_text_with_encoding(str(p))


def _extract_pdf(p: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(p))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        logger.warning("[file_parsers] PDF error for %s: %s", p, e)
        return f"[pdf error: {e}]"


def _extract_docx(p: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(p))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        logger.warning("[file_parsers] DOCX error for %s: %s", p, e)
        return f"[docx error: {e}]"


def _extract_doc(p: Path) -> str:
    """Extract text from legacy .doc files.

    Tries python-docx first (works for some .doc files saved in newer format),
    then falls back to reading raw bytes and decoding printable text.
    """
    try:
        from docx import Document
        doc = Document(str(p))
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception:
        pass
    # Raw byte extraction: grab printable ASCII / UTF-8 runs
    try:
        with open(p, "rb") as f:
            raw = f.read(MAX_INDEX_BYTES)
        text = raw.decode("utf-8", errors="replace")
        # Remove non-printable chars except newline/tab;
        # \u0400-\u04ff covers the Cyrillic Unicode block
        text = re.sub(r'[^\x09\x0a\x0d\x20-\x7e\u0400-\u04ff]', ' ', text)
        text = re.sub(r' {3,}', ' ', text)
        return text.strip()
    except Exception as e:
        logger.warning("[file_parsers] DOC error for %s: %s", p, e)
        return f"[doc error: {e}]"


def _extract_odt(p: Path) -> str:
    """Extract text from OpenDocument Text (.odt) files."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(str(p), "r") as z:
            if "content.xml" not in z.namelist():
                return ""
            content_xml = z.read("content.xml").decode("utf-8", errors="replace")
        root = ET.fromstring(content_xml)
        texts = []
        for elem in root.iter():
            if elem.text and elem.text.strip():
                texts.append(elem.text.strip())
            if elem.tail and elem.tail.strip():
                texts.append(elem.tail.strip())
        return "\n".join(texts)
    except Exception as e:
        logger.warning("[file_parsers] ODT error for %s: %s", p, e)
        return f"[odt error: {e}]"


def _extract_rtf(p: Path) -> str:
    """Extract plain text from RTF files via striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
        text = _read_text_with_encoding(str(p))
        return rtf_to_text(text)
    except ImportError:
        logger.warning("[file_parsers] striprtf not installed; reading RTF as plain text")
        return _read_text_with_encoding(str(p))
    except Exception as e:
        logger.warning("[file_parsers] RTF error for %s: %s", p, e)
        return f"[rtf error: {e}]"
